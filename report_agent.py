"""
Report Agent - Autonomous Report Generator
Collects outputs from all agents and generates comprehensive
Data Science reports in JSON, DOCX, and PDF formats.
"""

import json
import logging
import os
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Optional

from base_agent import BaseAgent

logging.basicConfig(level=logging.INFO, format="%(asctime)s [ReportAgent] %(message)s")
logger = logging.getLogger("ReportAgent")


class ReportAgent(BaseAgent):

    def __init__(self):
        super().__init__()
        self.collected_results: Dict[str, Any] = {}

    # ------------------------------------------------------------------
    # ✅ Capabilities
    # ------------------------------------------------------------------
    def get_capabilities(self) -> List[Dict[str, Any]]:
        return [
            {
                "function_name": "generate_dataset_summary",
                "description": "Generate executive summary section from dataset metadata",
                "parameters": ["dataset_name", "rows", "columns", "objective"],
                "examples": ["generate dataset summary", "create executive summary"]
            },
            {
                "function_name": "generate_eda_report",
                "description": "Generate EDA section from statistical summaries and correlations",
                "parameters": ["eda_results"],
                "examples": ["generate eda report", "summarize exploratory analysis"]
            },
            {
                "function_name": "generate_ml_report",
                "description": "Generate ML results section from model training outputs",
                "parameters": ["ml_results"],
                "examples": ["generate ml report", "summarize model results"]
            },
            {
                "function_name": "generate_feature_importance_report",
                "description": "Generate feature importance section",
                "parameters": ["feature_results"],
                "examples": ["generate feature importance report"]
            },
            {
                "function_name": "generate_recommendations",
                "description": "Generate recommendations section from CriticAgent output",
                "parameters": ["critic_results"],
                "examples": ["generate recommendations", "what should we improve"]
            },
            {
                "function_name": "export_json",
                "description": "Export full report as JSON file",
                "parameters": ["report_data", "filename"],
                "examples": ["export report as json"]
            },
            {
                "function_name": "export_docx",
                "description": "Export full report as DOCX file",
                "parameters": ["report_data", "filename"],
                "examples": ["export report as docx", "download word document"]
            },
            {
                "function_name": "export_pdf",
                "description": "Export full report as PDF file",
                "parameters": ["report_data", "filename"],
                "examples": ["export report as pdf", "download pdf report"]
            },
        ]

    # ------------------------------------------------------------------
    # ✅ Collect & Build
    # ------------------------------------------------------------------
    def collect_results(self, results: List[Dict[str, Any]]) -> None:
        """Collect all task results from orchestrator pipeline."""
        logger.info("Collecting results from pipeline")
        for item in results:
            task = item.get("task", "")
            result = item.get("result")
            if task and result is not None:
                self.collected_results[task] = result

    def build_report(
        self,
        dataset_name: str = "Dataset",
        objective: str = "Data Analysis & Machine Learning",
        rows: int = 0,
        columns: int = 0,
    ) -> Dict[str, Any]:
        """Build the full structured report from collected results."""
        logger.info("Generating report")

        report = {
            "metadata": {
                "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "dataset_name": dataset_name,
                "objective": objective,
            },
            "sections": {
                "executive_summary":        self.generate_dataset_summary(dataset_name, rows, columns, objective),
                "data_quality":             self._build_data_quality_section(),
                "eda":                      self._build_eda_section(),
                "visualizations":           self._build_visualization_section(),
                "ml_results":               self._build_ml_section(),
                "reflection_analysis":      self._build_reflection_section(),
                "feature_importance":       self._build_feature_importance_section(),
                "recommendations":          self._build_recommendations_section(),
                "conclusion":               self._build_conclusion_section(),
            }
        }
        logger.info("Report generation complete")
        return report

    # ------------------------------------------------------------------
    # ✅ Section Builders
    # ------------------------------------------------------------------
    def generate_dataset_summary(
        self,
        dataset_name: str = "Dataset",
        rows: int = 0,
        columns: int = 0,
        objective: str = "Data Analysis",
    ) -> Dict[str, Any]:
        ml = self._get_best_ml_result()
        key_findings = []

        if ml:
            acc = ml.get("accuracy") or ml.get("r2_score")
            if acc:
                key_findings.append(f"Best model '{ml.get('model', 'N/A')}' achieved {acc:.2%} accuracy")

        critic = self.collected_results.get("critic_evaluation", {})
        if isinstance(critic, dict):
            for issue in critic.get("issues_detected", [])[:2]:
                key_findings.append(issue)

        reflection = self.collected_results.get("reflection_loop", {})
        if isinstance(reflection, dict) and reflection.get("reflection_cycles", 0) > 0:
            key_findings.append(
                f"Reflection loop ran {reflection['reflection_cycles']} cycle(s), "
                f"improving accuracy to {reflection.get('best_accuracy', 0):.2%}"
            )

        return {
            "dataset_name": dataset_name,
            "rows":         rows,
            "columns":      columns,
            "objective":    objective,
            "key_findings": key_findings or ["Analysis completed successfully."],
        }

    def generate_eda_report(self, eda_results: Any = None) -> Dict[str, Any]:
        return self._build_eda_section()

    def generate_ml_report(self, ml_results: Any = None) -> Dict[str, Any]:
        return self._build_ml_section()

    def generate_feature_importance_report(self, feature_results: Any = None) -> Dict[str, Any]:
        return self._build_feature_importance_section()

    def generate_recommendations(self, critic_results: Any = None) -> Dict[str, Any]:
        return self._build_recommendations_section()

    # ------------------------------------------------------------------
    # ✅ Internal Section Helpers
    # ------------------------------------------------------------------
    def _build_data_quality_section(self) -> Dict[str, Any]:
        missing = {}
        duplicates_removed = "N/A"
        cleaning_actions = []

        eda_info = self.collected_results.get("get_data_info", {})
        if isinstance(eda_info, dict):
            missing = {k: v for k, v in eda_info.get("Missing Values", {}).items() if v > 0}

        for task in ("handle_missing", "handle_outliers", "remove_duplicates",
                     "encode_categoricals", "scale_features", "normalize_features"):
            if task in self.collected_results:
                cleaning_actions.append(task.replace("_", " ").title())

        return {
            "missing_values":      missing,
            "duplicates_removed":  duplicates_removed,
            "cleaning_actions":    cleaning_actions or ["No cleaning actions recorded"],
        }

    def _build_eda_section(self) -> Dict[str, Any]:
        stats = {}
        correlations = {}
        observations = []

        summary = self.collected_results.get("summary_statistics")
        if summary is not None:
            try:
                import pandas as pd
                if isinstance(summary, pd.DataFrame):
                    stats = summary.to_dict()
            except Exception:
                pass

        corr = self.collected_results.get("get_pairwise_correlation")
        if corr:
            correlations["pairwise"] = str(corr)

        missing = self.collected_results.get("missing_values")
        if missing is not None:
            observations.append("Missing value analysis completed")

        return {
            "statistical_summary": stats,
            "correlations":        correlations,
            "observations":        observations or ["EDA completed."],
        }

    def _build_visualization_section(self) -> Dict[str, Any]:
        charts = []
        viz_tasks = ("plot_histogram", "plot_boxplot", "plot_scatter",
                     "plot_correlation_heatmap", "correlation_heatmap", "plot_bar_chart")
        for task in viz_tasks:
            if task in self.collected_results:
                charts.append({
                    "chart_type":  task.replace("_", " ").title(),
                    "description": f"{task.replace('_', ' ').title()} generated during analysis"
                })
        return {"charts_generated": charts or ["No visualizations recorded"]}

    def _build_ml_section(self) -> Dict[str, Any]:
        algorithms = []
        best_model = {}

        for task in ("train_classification", "train_regression"):
            result = self.collected_results.get(task)
            if isinstance(result, dict):
                algorithms.append({
                    "model":    result.get("model", "N/A"),
                    "accuracy": result.get("accuracy"),
                    "rmse":     result.get("rmse"),
                    "r2_score": result.get("r2_score"),
                    "target":   result.get("target", "N/A"),
                })
                best_model = result

        reflection = self.collected_results.get("reflection_loop", {})
        if isinstance(reflection, dict) and reflection.get("best_accuracy"):
            best_model = reflection.get("final_result", best_model)

        return {
            "algorithms_tested": algorithms,
            "best_model":        best_model.get("model", "N/A"),
            "best_accuracy":     best_model.get("accuracy") or best_model.get("r2_score"),
            "classification_report": best_model.get("classification_report", ""),
            "confusion_matrix":  best_model.get("confusion_matrix", []),
        }

    def _build_reflection_section(self) -> Dict[str, Any]:
        reflection = self.collected_results.get("reflection_loop", {})
        if not isinstance(reflection, dict):
            return {"status": "Reflection loop not executed"}

        return {
            "cycles_executed":     reflection.get("reflection_cycles", 0),
            "improvements_applied": reflection.get("improvements_applied", []),
            "best_accuracy":       reflection.get("best_accuracy"),
            "best_model":          reflection.get("best_model"),
            "history":             reflection.get("reflection_history", []),
        }

    def _build_feature_importance_section(self) -> Dict[str, Any]:
        fi = self.collected_results.get("feature_importance")
        if fi is None:
            return {"status": "Feature importance not computed"}

        try:
            import pandas as pd
            if isinstance(fi, pd.DataFrame):
                top = fi.head(10)
                return {
                    "top_features": top.to_dict(orient="records"),
                    "note": "Higher score = more influential feature"
                }
        except Exception:
            pass

        return {"feature_importance_data": str(fi)}

    def _build_recommendations_section(self) -> Dict[str, Any]:
        critic = self.collected_results.get("critic_evaluation", {})
        recs = []
        issues = []

        if isinstance(critic, dict):
            recs   = critic.get("recommendations", [])
            issues = critic.get("issues_detected", [])

        next_actions = [
            "Collect more labeled data to improve model generalization",
            "Monitor model performance in production with periodic retraining",
            "Consider deploying best model via REST API for real-time inference",
        ]

        return {
            "model_improvements":    recs or ["No specific improvements flagged"],
            "issues_detected":       issues,
            "suggested_next_actions": next_actions,
        }

    def _build_conclusion_section(self) -> Dict[str, Any]:
        ml_section = self._build_ml_section()
        critic = self.collected_results.get("critic_evaluation", {})
        severity = critic.get("severity", "low") if isinstance(critic, dict) else "low"

        confidence_map = {"low": "High", "medium": "Medium", "high": "Low"}
        confidence = confidence_map.get(severity, "Medium")

        return {
            "best_model":       ml_section.get("best_model", "N/A"),
            "best_accuracy":    ml_section.get("best_accuracy"),
            "confidence_level": confidence,
            "summary":          (
                f"Analysis completed. Best model: {ml_section.get('best_model', 'N/A')}. "
                f"Overall confidence: {confidence}."
            )
        }

    def _get_best_ml_result(self) -> Optional[Dict]:
        reflection = self.collected_results.get("reflection_loop", {})
        if isinstance(reflection, dict) and reflection.get("final_result"):
            return reflection["final_result"]
        return (
            self.collected_results.get("train_classification") or
            self.collected_results.get("train_regression")
        )

    # ------------------------------------------------------------------
    # ✅ Export: JSON
    # ------------------------------------------------------------------
    def export_json(
        self,
        report_data: Dict[str, Any],
        filename: str = "report.json"
    ) -> Dict[str, Any]:
        logger.info("Exporting JSON report")
        try:
            path = os.path.join(tempfile.gettempdir(), filename)
            with open(path, "w", encoding="utf-8") as f:
                json.dump(report_data, f, indent=2, default=str)
            logger.info(f"JSON report saved: {path}")
            return {"report_generated": True, "file_type": "json", "file_path": path, "sections": list(report_data.get("sections", {}).keys())}
        except Exception as e:
            return {"report_generated": False, "error": str(e)}

    # ------------------------------------------------------------------
    # ✅ Export: DOCX
    # ------------------------------------------------------------------
    def export_docx(
        self,
        report_data: Dict[str, Any],
        filename: str = "report.docx"
    ) -> Dict[str, Any]:
        logger.info("Exporting DOCX report")
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            return {"report_generated": False, "error": "python-docx not installed. Run: pip install python-docx"}

        try:
            doc = Document()
            meta = report_data.get("metadata", {})

            # Title
            title = doc.add_heading("Data Science Report", 0)
            title.runs[0].font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

            doc.add_paragraph(f"Generated: {meta.get('generated_at', '')}")
            doc.add_paragraph(f"Dataset: {meta.get('dataset_name', '')}")
            doc.add_paragraph(f"Objective: {meta.get('objective', '')}")

            sections = report_data.get("sections", {})
            section_titles = {
                "executive_summary":   "Section 1: Executive Summary",
                "data_quality":        "Section 2: Data Quality Report",
                "eda":                 "Section 3: Exploratory Data Analysis",
                "visualizations":      "Section 4: Visualizations",
                "ml_results":          "Section 5: Machine Learning Results",
                "reflection_analysis": "Section 6: Reflection Analysis",
                "feature_importance":  "Section 7: Feature Importance",
                "recommendations":     "Section 8: Recommendations",
                "conclusion":          "Section 9: Conclusion",
            }

            for key, heading in section_titles.items():
                data = sections.get(key, {})
                doc.add_heading(heading, level=1)
                self._write_dict_to_doc(doc, data)

            path = os.path.join(tempfile.gettempdir(), filename)
            doc.save(path)
            logger.info(f"DOCX report saved: {path}")
            return {"report_generated": True, "file_type": "docx", "file_path": path, "sections": list(sections.keys())}
        except Exception as e:
            return {"report_generated": False, "error": str(e)}

    def _write_dict_to_doc(self, doc, data: Any, indent: int = 0) -> None:
        """Recursively write dict/list content to Word document."""
        prefix = "  " * indent
        if isinstance(data, dict):
            for k, v in data.items():
                if isinstance(v, (dict, list)):
                    doc.add_paragraph(f"{prefix}{k.replace('_', ' ').title()}:")
                    self._write_dict_to_doc(doc, v, indent + 1)
                else:
                    doc.add_paragraph(f"{prefix}{k.replace('_', ' ').title()}: {str(v)[:500]}")
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, (dict, list)):
                    self._write_dict_to_doc(doc, item, indent + 1)
                else:
                    doc.add_paragraph(f"{prefix}\u2022 {str(item)[:500]}")
        else:
            doc.add_paragraph(f"{prefix}{str(data)[:500]}")

    # ------------------------------------------------------------------
    # ✅ Export: PDF
    # ------------------------------------------------------------------
    def export_pdf(
        self,
        report_data: Dict[str, Any],
        filename: str = "report.pdf"
    ) -> Dict[str, Any]:
        logger.info("Exporting PDF report")
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        except ImportError:
            return {"report_generated": False, "error": "reportlab not installed. Run: pip install reportlab"}

        try:
            path = os.path.join(tempfile.gettempdir(), filename)
            doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            accent = ParagraphStyle("Accent", parent=styles["Heading1"],
                                    textColor=colors.HexColor("#667eea"), fontSize=14)
            body   = styles["Normal"]
            story  = []

            meta = report_data.get("metadata", {})
            story.append(Paragraph("Data Science Report", styles["Title"]))
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph(f"Generated: {meta.get('generated_at', '')}", body))
            story.append(Paragraph(f"Dataset: {meta.get('dataset_name', '')}", body))
            story.append(Paragraph(f"Objective: {meta.get('objective', '')}", body))
            story.append(Spacer(1, 0.5*cm))

            sections = report_data.get("sections", {})
            section_titles = {
                "executive_summary":   "Section 1: Executive Summary",
                "data_quality":        "Section 2: Data Quality Report",
                "eda":                 "Section 3: Exploratory Data Analysis",
                "visualizations":      "Section 4: Visualizations",
                "ml_results":          "Section 5: Machine Learning Results",
                "reflection_analysis": "Section 6: Reflection Analysis",
                "feature_importance":  "Section 7: Feature Importance",
                "recommendations":     "Section 8: Recommendations",
                "conclusion":          "Section 9: Conclusion",
            }

            for key, heading in section_titles.items():
                data = sections.get(key, {})
                story.append(HRFlowable(width="100%", color=colors.HexColor("#667eea")))
                story.append(Paragraph(heading, accent))
                story.append(Spacer(1, 0.2*cm))
                self._write_dict_to_pdf(story, data, body)
                story.append(Spacer(1, 0.4*cm))

            doc.build(story)
            logger.info(f"PDF report saved: {path}")
            return {"report_generated": True, "file_type": "pdf", "file_path": path, "sections": list(sections.keys())}
        except Exception as e:
            return {"report_generated": False, "error": str(e)}

    def _write_dict_to_pdf(self, story: list, data: Any, style, indent: int = 0) -> None:
        """Recursively write dict/list content into PDF story."""
        from reportlab.platypus import Paragraph
        import html
        prefix = "&nbsp;" * (indent * 4)
        if isinstance(data, dict):
            for k, v in data.items():
                if isinstance(v, (dict, list)):
                    story.append(Paragraph(f"{prefix}<b>{html.escape(k.replace('_', ' ').title())}:</b>", style))
                    self._write_dict_to_pdf(story, v, style, indent + 1)
                else:
                    safe_v = html.escape(str(v))[:300]
                    story.append(Paragraph(f"{prefix}<b>{html.escape(k.replace('_', ' ').title())}:</b> {safe_v}", style))
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, (dict, list)):
                    self._write_dict_to_pdf(story, item, style, indent + 1)
                else:
                    story.append(Paragraph(f"{prefix}• {html.escape(str(item))[:300]}", style))
        else:
            story.append(Paragraph(f"{prefix}{html.escape(str(data))[:300]}", style))

    # ------------------------------------------------------------------
    # ✅ Safe dispatcher
    # ------------------------------------------------------------------
    def execute_capability(self, function_name: str, **kwargs) -> Any:
        if not hasattr(self, function_name):
            return f"Method '{function_name}' not found in ReportAgent."
        try:
            method = getattr(self, function_name)
            return method(**kwargs) if callable(method) else method
        except Exception as e:
            return f"❌ Error executing {function_name} in ReportAgent: {e}"
